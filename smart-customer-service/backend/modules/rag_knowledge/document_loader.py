"""
文档加载器模块 - 支持多格式文档解析和智能分块。

提供PDF、Word、Markdown、HTML等格式的文档解析功能，
以及多种分块策略（递归字符、段落、标题层级）。
"""

import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langfuse import observe

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据模型。

    Attributes:
        page_content: 文档内容文本
        metadata: 元数据字典，包含来源、页码、标题等信息
        doc_id: 文档唯一标识符（UUID）
        embedding: 可选的向量嵌入
    """

    page_content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = ""
    embedding: Optional[List[float]] = None

    def __post_init__(self):
        """初始化后生成doc_id（如果未提供）。"""
        if not self.doc_id:
            content_hash = hashlib.md5(self.page_content.encode()).hexdigest()[:16]
            self.doc_id = f"doc_{content_hash}"


class DocumentParser:
    """文档解析器 - 支持多种文档格式。

    支持PDF、DOCX、Markdown、HTML格式的文档解析。
    """

    SUPPORTED_FORMATS = {".pdf", ".docx", ".md", ".markdown", ".html", ".htm"}

    @observe(name="parse_document", as_type="span")
    def parse(self, file_path: Path) -> str:
        """解析文档文件，提取文本内容。

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内容

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
            Exception: 解析过程中的其他错误
        """
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {ext}。支持的格式: {self.SUPPORTED_FORMATS}")

        logger.info(f"开始解析文档: {file_path}")

        try:
            if ext == ".pdf":
                text = self._parse_pdf(file_path)
            elif ext == ".docx":
                text = self._parse_docx(file_path)
            elif ext in [".md", ".markdown"]:
                text = self._parse_markdown(file_path)
            elif ext in [".html", ".htm"]:
                text = self._parse_html(file_path)
            else:
                raise ValueError(f"不支持的格式: {ext}")

            logger.info(f"文档解析完成，提取 {len(text)} 字符")
            return text

        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            raise

    def _parse_pdf(self, file_path: Path) -> str:
        """解析PDF文件。

        Args:
            file_path: PDF文件路径

        Returns:
            提取的文本内容
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            text_parts = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)

            doc.close()
            return "\n\n".join(text_parts)

        except ImportError:
            logger.warning("PyMuPDF未安装，尝试使用备用方法")
            return self._parse_pdf_fallback(file_path)

    def _parse_pdf_fallback(self, file_path: Path) -> str:
        """PDF解析备用方法（使用PyPDF2）。

        Args:
            file_path: PDF文件路径

        Returns:
            提取的文本内容
        """
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except ImportError:
            raise ImportError(
                "需要安装 PyMuPDF 或 PyPDF2 来解析PDF文件。"
                "请运行: pip install PyMuPDF 或 pip install PyPDF2"
            )

    def _parse_docx(self, file_path: Path) -> str:
        """解析Word文档。

        Args:
            file_path: DOCX文件路径

        Returns:
            提取的文本内容
        """
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)

        except ImportError:
            raise ImportError(
                "需要安装 python-docx 来解析Word文档。"
                "请运行: pip install python-docx"
            )

    def _parse_markdown(self, file_path: Path) -> str:
        """解析Markdown文件。

        Args:
            file_path: Markdown文件路径

        Returns:
            提取的文本内容
        """
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_html(self, file_path: Path) -> str:
        """解析HTML文件。

        Args:
            file_path: HTML文件路径

        Returns:
            提取的文本内容
        """
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                return soup.get_text(separator="\n", strip=True)

        except ImportError:
            raise ImportError(
                "需要安装 beautifulsoup4 来解析HTML文件。"
                "请运行: pip install beautifulsoup4"
            )


class DocumentChunker:
    """文档分块器 - 支持多种分块策略。

    支持递归字符分割、段落分割、标题层级分割等策略。
    """

    def __init__(
        self,
        strategy: str = "recursive",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        """初始化分块器。

        Args:
            strategy: 分块策略 ('recursive', 'paragraph', 'title')
            chunk_size: 每个块的大小（字符数）
            chunk_overlap: 块之间的重叠字符数
        """
        self.strategy = strategy
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    @observe(name="chunk_documents", as_type="span")
    def chunk(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """将文本分块为多个文档。

        Args:
            text: 要分块的文本内容
            metadata: 元数据字典，会添加到每个分块中

        Returns:
            分块后的文档列表
        """
        logger.info(f"开始分块，策略: {self.strategy}, 块大小: {self.chunk_size}")

        if self.strategy == "recursive":
            chunks = self._recursive_chunk(text)
        elif self.strategy == "paragraph":
            chunks = self._paragraph_chunk(text)
        elif self.strategy == "title":
            chunks = self._title_chunk(text)
        else:
            raise ValueError(f"不支持的分块策略: {self.strategy}")

        documents = []
        for i, chunk_text in enumerate(chunks):
            if chunk_text.strip():
                chunk_metadata = metadata.copy() if metadata else {}
                chunk_metadata["chunk_index"] = i
                chunk_metadata["chunk_strategy"] = self.strategy

                doc = Document(
                    page_content=chunk_text.strip(),
                    metadata=chunk_metadata,
                )
                documents.append(doc)

        logger.info(f"分块完成，共 {len(documents)} 个块")
        return documents

    def _recursive_chunk(self, text: str) -> List[str]:
        """使用递归字符分割器分块。

        Args:
            text: 要分块的文本

        Returns:
            分块后的文本列表
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )
        return splitter.split_text(text)

    def _paragraph_chunk(self, text: str) -> List[str]:
        """按段落分块。

        Args:
            text: 要分块的文本

        Returns:
            分块后的文本列表
        """
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = para

        if current_chunk:
            chunks.append(current_chunk)

        return chunks

    def _title_chunk(self, text: str) -> List[str]:
        """按标题层级分块。

        Args:
            text: 要分块的文本

        Returns:
            分块后的文本列表
        """
        import re

        # 匹配Markdown标题 (# Title, ## Subtitle, etc.)
        title_pattern = r"^(#{1,6})\s+(.+)$"
        lines = text.split("\n")

        chunks = []
        current_chunk = ""
        current_title = ""

        for line in lines:
            match = re.match(title_pattern, line)
            if match:
                # 遇到新标题，保存当前块
                if current_chunk:
                    chunks.append(current_chunk)
                current_title = line
                current_chunk = line + "\n"
            else:
                if len(current_chunk) + len(line) <= self.chunk_size:
                    current_chunk += line + "\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = line + "\n"

        if current_chunk:
            chunks.append(current_chunk)

        return chunks


@observe(name="import_documents", as_type="span")
def import_documents(
    file_paths: List[Path],
    chunk_strategy: str = "recursive",
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    metadata: Optional[Dict[str, Any]] = None,
) -> List[Document]:
    """导入并处理多个文档文件。

    Args:
        file_paths: 文档文件路径列表
        chunk_strategy: 分块策略
        chunk_size: 块大小
        chunk_overlap: 块重叠大小
        metadata: 额外元数据

    Returns:
        处理后的文档列表
    """
    parser = DocumentParser()
    chunker = DocumentChunker(
        strategy=chunk_strategy,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    all_documents = []

    for file_path in file_paths:
        logger.info(f"处理文件: {file_path}")

        # 解析文档
        text = parser.parse(file_path)

        # 构建基础元数据
        file_metadata = metadata.copy() if metadata else {}
        file_metadata["source"] = str(file_path)
        file_metadata["filename"] = file_path.name

        # 分块
        documents = chunker.chunk(text, metadata=file_metadata)
        all_documents.extend(documents)

    logger.info(f"文档导入完成，共 {len(all_documents)} 个文档块")
    return all_documents
